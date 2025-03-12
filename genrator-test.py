import json
import gradio as gr
from docx import Document
import tempfile
import os
import openai
import csv
from openai import OpenAI
from dotenv import load_dotenv
import requests
from bs4 import BeautifulSoup
import re

# Load API key from .env file
load_dotenv()

# Initialize OpenAI client
client = OpenAI()

def generate_test_cases(requirements_text, num_test_cases, output_format):
    try:
        # Structured prompt with exact format
        prompt = f"""Generate {num_test_cases} detailed test cases for the following software requirements. Format each test case EXACTLY as shown in the example below, with no deviations in formatting or structure.

Requirements:
{requirements_text}

    Generate the test cases in this exact format:

    Test Case 1: [Title]
    Test Case ID: TC-001
    Description: [Detailed description of the test case]
    Preconditions: [List of required preconditions]
    Test Steps:

    [Numbered list of steps]
    Expected Results: [Detailed expected outcome]
    Test Data:
    [Key-value pairs of test data]

    Test Case 2: [Title]
    [Continue same format]

    Example of exact formatting:
    Test Case 1: Successful Login
    Test Case ID: TC-001
    Description: Verify that a user can successfully log in with a valid email and password.
    Preconditions: A registered user exists with valid credentials.
    Test Steps:

    1. Navigate to the login page.
    2. Enter a valid email and password.
    3. Click the "Login" button.
    Expected Results: The user is successfully logged in and redirected to the dashboard.
    Test Data:
    Email: testuser@example.com
    Password: SecurePass123

Please generate exactly {num_test_cases} test cases following this exact format."""

        # Generate response using OpenAI API
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are a QA expert who creates detailed test cases. Generate exactly {num_test_cases} test cases with the exact formatting requested."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )

        # Extract and clean the generated text
        formatted_response = response.choices[0].message.content
        formatted_response = formatted_response.replace("```", "").strip()
        formatted_response = formatted_response.replace("\n\n\n", "\n\n")
        
        # Save based on selected format and return both text and file
        if output_format == "DOCX":
            return formatted_response, save_as_docx(formatted_response, "test_cases")
        elif output_format == "CSV":
            return formatted_response, save_as_csv(formatted_response, "test_cases")
        else:
            return formatted_response, save_as_txt(formatted_response, "test_cases")

    except Exception as e:
        error_msg = f"Error generating test cases: {str(e)}"
        temp_dir = tempfile.gettempdir()
        error_file = os.path.join(temp_dir, "error.txt")
        with open(error_file, 'w') as f:
            f.write(error_msg)
        return error_msg, error_file

def format_test_cases(raw_text):
    try:
        # Remove any extra newlines
        formatted_text = raw_text.replace("\n\n\n", "\n\n")
        
        # Ensure consistent formatting
        formatted_text = formatted_text.replace("Test Case ", "\n\nTest Case ")
        
        # Add proper spacing for sections
        sections = ["Test Case ID:", "Description:", "Preconditions:", 
                   "Test Steps:", "Expected Results:", "Test Data:"]
        
        for section in sections:
            formatted_text = formatted_text.replace(section, f"\n{section}")
        
        # Add header
        formatted_text = "TEST CASES\n" + "="*50 + "\n\n" + formatted_text.strip()
        
        return formatted_text
    except Exception as e:
        return raw_text

def save_as_docx(text, file_name="output"):
    try:
        doc = Document()
        doc.add_heading('Test Cases', 0)
        
        # Split test cases and format them
        test_cases = text.split("\n\nTest Case ")
        
        for idx, tc in enumerate(test_cases):
            if idx == 0 and not tc.strip():
                continue
                
            if idx > 0:
                tc = "Test Case " + tc
            
            # Add each section with proper formatting
            sections = tc.split("\n")
            for section in sections:
                if section.strip():
                    if section.startswith("Test Case "):
                        doc.add_heading(section, level=2)
                    elif any(s in section for s in ["Test Case ID:", "Description:", 
                                                  "Preconditions:", "Test Steps:", 
                                                  "Expected Results:", "Test Data:"]):
                        p = doc.add_paragraph()
                        p.add_run(section).bold = True
                    else:
                        doc.add_paragraph(section)
            
            # Add spacing between test cases
            doc.add_paragraph()
        
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, f"{file_name}.docx")
        doc.save(file_path)
        return file_path
    except Exception as e:
        temp_dir = tempfile.gettempdir()
        error_file = os.path.join(temp_dir, "error.txt")
        with open(error_file, 'w') as f:
            f.write(f"Error generating DOCX: {str(e)}")
        return error_file

def save_as_txt(text, file_name="output"):
    try:
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, f"{file_name}.txt")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return file_path
    except Exception as e:
        temp_dir = tempfile.gettempdir()
        error_file = os.path.join(temp_dir, "error.txt")
        with open(error_file, 'w') as f:
            f.write(f"Error saving text: {str(e)}")
        return error_file

def save_as_csv(text, file_name="output"):
    try:
        test_cases = parse_test_cases(text)
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, f"{file_name}.csv")
        
        # Get all unique fields
        fields = set()
        for tc in test_cases:
            fields.update(tc.keys())
        fields = ['Title', 'Test Case ID', 'Description', 'Preconditions', 
                 'Test Steps', 'Expected Results', 'Test Data'] + \
                [f for f in fields if f not in ['Title', 'Test Case ID', 'Description', 
                                              'Preconditions', 'Test Steps', 'Expected Results', 'Test Data']]
        
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=fields)
            writer.writeheader()
            writer.writerows(test_cases)
            
        return file_path
    except Exception as e:
        return f"Error saving CSV: {str(e)}"

def parse_test_cases(text):
    """Parse the generated text into structured test cases"""
    test_cases = []
    current_case = {}
    current_field = None
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('Test Case ') and ':' in line:
            if current_case:
                test_cases.append(current_case)
            current_case = {'Title': line.split(':', 1)[1].strip()}
        elif ':' in line:
            field, value = line.split(':', 1)
            field = field.strip()
            value = value.strip()
            current_case[field] = value
            current_field = field
        elif current_field:
            current_case[current_field] = current_case.get(current_field, '') + '\n' + line
            
    if current_case:
        test_cases.append(current_case)
        
    return test_cases

def analyze_requirements(user_stories, acceptance_criteria):
    try:
        prompt = f"""Analyze the following user stories and acceptance criteria. 
        Provide a clear summary, identify any gaps or potential edge cases, and suggest additional considerations.

        User Stories:
        {user_stories}

        Acceptance Criteria:
        {acceptance_criteria}

        Please format your response exactly as follows:
        
        📝 Summary:
        [Provide a clear, concise summary of the requirements]

        🔍 Analysis:
        [Analyze the completeness and clarity of the requirements]

        🚧 Identified Gaps:
        1. [Gap or missing requirement]
        2. [Gap or missing requirement]
        ...

        💡 Additional Test Scenarios to Consider:
        1. [Additional test scenario]
        2. [Additional test scenario]
        ...
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a requirements analysis expert. Analyze the given requirements thoroughly and identify potential gaps and edge cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        return response.choices[0].message.content
    except Exception as e:
        return f"Error analyzing requirements: {str(e)}"

def extract_requirements_from_url(url):
    try:
        # Fetch the webpage content
        response = requests.get(url)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Initialize lists to store found content
        user_stories = []
        acceptance_criteria = []
        
        # Search through HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
            text = element.get_text().strip()
            
            # Look for User Story sections
            if re.search(r'user\s+stor(y|ies)', text.lower()):
                next_elem = element.find_next_sibling()
                while next_elem and not next_elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    if next_elem.get_text().strip():
                        user_stories.append(next_elem.get_text().strip())
                    next_elem = next_elem.find_next_sibling()
            
            # Look for Acceptance Criteria sections
            if re.search(r'acceptance\s+criteria', text.lower()):
                next_elem = element.find_next_sibling()
                while next_elem and not next_elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    if next_elem.get_text().strip():
                        acceptance_criteria.append(next_elem.get_text().strip())
                    next_elem = next_elem.find_next_sibling()
        
        # Get AI analysis of the requirements
        analysis = analyze_requirements("\n".join(user_stories), "\n".join(acceptance_criteria))
        
        # Format the output
        formatted_output = "🔎 Requirements Analysis\n"
        formatted_output += "=" * 50 + "\n\n"
        formatted_output += analysis
        
        return formatted_output
    
    except Exception as e:
        return f"Error extracting requirements from URL: {str(e)}"

def create_interface():
    with gr.Blocks(css="""
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap');
        
        .gradio-container {
            background: linear-gradient(135deg, #e6e6fa 0%, #b39ddb 100%);
            min-height: 100vh;
        }
        .main-header {
            background: linear-gradient(90deg, #9575cd 0%, #7e57c2 100%);
            color: white;
            padding: 1.2rem;
            border-radius: 10px;
            margin-bottom: 1.5rem;
            text-align: center;
        }
        .main-header h1 {
            color: white !important;
            font-size: 2.8rem !important;
            font-weight: 700 !important;
            font-family: 'Roboto', sans-serif !important;
            margin: 0 !important;
        }
        .sub-header {
            background: rgba(255, 255, 255, 0.8);
            padding: 1rem;
            border-radius: 8px;
            margin-bottom: 1.5rem;
            color: #5e35b1;
            font-size: 1.5rem !important;
            text-align: center;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
            font-family: 'Roboto', sans-serif !important;
            font-weight: 500 !important;
        }
        .input-panel {
            background: rgba(255, 255, 255, 0.9);
            padding: 2rem;
            border-radius: 12px;
            box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
            margin: 0.5rem;
            min-height: 550px;
            width: 110% !important;
        }
        .output-panel {
            background: rgba(255, 255, 255, 0.9);
            padding: 1.5rem;
            border-radius: 12px;
            box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
            margin: 0.5rem;
            min-height: 550px;
        }
        .button-row {
            display: flex;
            gap: 1rem;
            margin-top: 2rem;
            justify-content: space-between;
            flex-wrap: wrap;
        }
        /* Base button styles */
        button.primary-button, button.secondary-button, button.clear-button {
            background: linear-gradient(90deg, #9575cd 0%, #7e57c2 100%) !important;
            color: white !important;
            transition: all 0.3s ease !important;
            border: none !important;
            box-shadow: 0 3px 7px rgba(0,0,0,0.2) !important;
            padding: 0.7rem 1.2rem !important;
            font-size: 1rem !important;
            font-family: 'Roboto', sans-serif !important;
            font-weight: 500 !important;
            min-width: 140px !important;
            text-align: center !important;
            white-space: nowrap !important;
            overflow: hidden !important;
            text-overflow: ellipsis !important;
        }
        /* Consistent hover effect for all buttons */
        button.primary-button:hover, button.secondary-button:hover, button.clear-button:hover {
            background: linear-gradient(90deg, #5e35b1 0%, #4527a0 100%) !important;
            transform: translateY(-3px) !important;
            box-shadow: 0 5px 10px rgba(0,0,0,0.3) !important;
        }
        /* Focus styles for accessibility */
        button:focus {
            outline: 2px solid #fff !important;
            outline-offset: 2px !important;
        }
        /* Make form elements larger */
        .input-panel label {
            font-size: 1.1rem !important;
            margin-bottom: 0.5rem !important;
            font-family: 'Roboto', sans-serif !important;
            font-weight: 500 !important;
        }
        .input-panel input, .input-panel textarea {
            font-size: 1.05rem !important;
            padding: 0.7rem !important;
            font-family: 'Roboto', sans-serif !important;
        }
        /* Remove inner borders */
        .gradio-container .gr-form, .gradio-container .gr-group, .gradio-container .gr-box {
            border: none !important;
            background: transparent !important;
        }
        .gradio-container .gr-input, .gradio-container .gr-textarea, .gradio-container .gr-radio {
            border-color: #b39ddb !important;
        }
        .gradio-container .gr-panel {
            border: none !important;
        }
        .gradio-container .gr-padded {
            padding: 0 !important;
        }
        /* Format code blocks in output */
        .output-panel pre {
            background-color: #f5f5f5 !important;
            border-radius: 5px !important;
            padding: 10px !important;
            border-left: 4px solid #9575cd !important;
            font-family: 'Consolas', 'Monaco', monospace !important;
            overflow-x: auto !important;
            margin: 10px 0 !important;
        }
        .output-panel code {
            font-family: 'Consolas', 'Monaco', monospace !important;
            background-color: #f5f5f5 !important;
            padding: 2px 4px !important;
            border-radius: 3px !important;
        }
        /* Change radio button color to violet */
        .gradio-container input[type="radio"]:checked {
            background-color: #7e57c2 !important;
            border-color: #7e57c2 !important;
        }
        .gradio-container input[type="radio"]:checked::before {
            background-color: #7e57c2 !important;
        }
        .gradio-container input[type="radio"]:focus {
            border-color: #7e57c2 !important;
            box-shadow: 0 0 0 2px rgba(126, 87, 194, 0.3) !important;
        }
    """) as iface:
        with gr.Column():
            gr.Markdown("# Personal Pookie Generator 🎀", elem_classes=["main-header"])
            gr.Markdown("""Generate detailed test cases, API documentation, and frontend validations from requirements.""", 
                       elem_classes=["sub-header"])
        
        with gr.Row():
            with gr.Column(scale=6, elem_classes=["input-panel"]):
                confluence_link_input = gr.Textbox(
                    label="Confluence Link", 
                    placeholder="Enter the Confluence link here...",
                    lines=2
                )
                num_cases = gr.Number(
                    label="Number of Test Cases",
                    value=5,
                    minimum=1,
                    maximum=20,
                    step=1
                )
                output_format = gr.Radio(
                    ["TXT", "DOCX", "CSV"],
                    label="Output Format", 
                    value="TXT"
                )
                
                with gr.Row(elem_classes=["button-row"]):
                    clear_btn = gr.Button("Clear", elem_classes=["clear-button"])
                    generate_test_btn = gr.Button("Test Cases", elem_classes=["primary-button"])
                    generate_api_btn = gr.Button("API Doc", elem_classes=["secondary-button"])
                    generate_fe_btn = gr.Button("FE Validation", elem_classes=["secondary-button"])
                
            with gr.Column(scale=7, elem_classes=["output-panel"]):
                output_text = gr.Textbox(
                    label="Generated Content",
                    lines=22,
                    interactive=False
                )
                output_file = gr.File(label="Download Generated File")
        
        # Button actions
        def clear_all():
            return [None, None, None]
        
        clear_btn.click(
            clear_all,
            inputs=None,
            outputs=[confluence_link_input, output_text, output_file]
        )
        
        generate_test_btn.click(
            generate_test_cases,
            inputs=[confluence_link_input, num_cases, output_format],
            outputs=[output_text, output_file]
        )
        
        generate_api_btn.click(
            generate_api_doc,
            inputs=[confluence_link_input, output_format],
            outputs=[output_text, output_file]
        )

        generate_fe_btn.click(
            generate_fe_validation,
            inputs=[confluence_link_input, output_format],
            outputs=[output_text, output_file]
        )
        
        return iface

def generate_api_doc(requirements_text, output_format):
    try:
        prompt = f"""Generate detailed Backend API documentation for the following requirements.
        Follow the EXACT format shown in the example below.

        Requirements:
        {requirements_text}

        Format each API endpoint EXACTLY like this:

        [Feature Name] API
        Endpoint:
        [HTTP_METHOD] [/api/path]

        Request Body:
        ```json
        {{
            "field1": "value1",
            "field2": "value2"
        }}
        ```

        Response:
        ✅ Success ([STATUS_CODE] [STATUS_TEXT])
        ```json
        {{
            "success": true,
            "message": "[Success message]",
            "data": {{
                "field": "value"
            }}
        }}
        ```

        ❌ Failure ([STATUS_CODE] [ERROR_TYPE])
        ```json
        {{
            "success": false,
            "message": "[Error message]"
        }}
        ```

        Example format for a login API:
        User Login API
        Endpoint:
        POST /api/auth/login

        Request Body:
        ```json
        {{
          "email": "testuser@example.com",
          "password": "SecurePass123"
        }}
        ```

        Response:
        ✅ Success (200 OK)
        ```json
        {{
          "success": true,
          "message": "Login successful",
          "token": "eyJhbGciOiJIUzI1NiIsInR5c..."
        }}
        ```

        ❌ Failure (401 Unauthorized)
        ```json
        {{
          "success": false,
          "message": "Invalid email or password"
        }}
        ```

        Generate the API documentation following this exact format for all endpoints needed to fulfill the requirements.
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a backend API documentation expert. Create detailed API documentation following the exact format provided."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )

        formatted_response = response.choices[0].message.content
        # Clean up the response but preserve code blocks
        formatted_response = formatted_response.strip()

        if output_format == "DOCX":
            return formatted_response, save_as_docx(formatted_response, "BE_documentation")
        elif output_format == "CSV":
            return formatted_response, save_as_csv(formatted_response, "BE_documentation")
        else:
            return formatted_response, save_as_txt(formatted_response, "BE_documentation")

    except Exception as e:
        error_msg = f"Error generating API documentation: {str(e)}"
        temp_dir = tempfile.gettempdir()
        error_file = os.path.join(temp_dir, "error.txt")
        with open(error_file, 'w') as f:
            f.write(error_msg)
        return error_msg, error_file

def generate_fe_validation(requirements_text, output_format):
    try:
        prompt = f"""Generate detailed Frontend validation rules and code for the following requirements.
        Follow the EXACT format shown in the example below.

        Requirements:
        {requirements_text}

        Format the validation documentation EXACTLY like this:

        🔹 [Form Name] Validation
        Validation Rules:
        [Field]: [Validation rules list]

        Example Validation Code (React with Formik & Yup):
        ```javascript
        import * as Yup from "yup";

        const validationSchema = Yup.object({{
            field: Yup.string()
                .required("[Field] is required")
                .validation1("[Error message]")
                .validation2("[Error message]")
        }});

        export default validationSchema;
        ```

        🔹 UI Enhancement Rules:
        1. [Enhancement rule]
        2. [Enhancement rule]

        🔹 Error Message Mapping:
        Field | Error Message
        --- | ---
        [Field] | "[Error message]"

        🔹 API Response Handling:
        ```javascript
        const handleAPIResponse = (response) => {{
            switch (response.status) {{
                case 200:
                    // Success handling
                    break;
                case 400:
                    // Error handling
                    break;
            }}
        }};
        ```

        Include for each form:
        1. Validation rules in plain text
        2. Yup validation schema
        3. UI enhancement suggestions
        4. Error message mapping
        5. API response handling
        """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a frontend validation expert. Create detailed validation rules and code following the exact format provided."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )

        formatted_response = response.choices[0].message.content
        formatted_response = formatted_response.strip()

        if output_format == "DOCX":
            return formatted_response, save_as_docx(formatted_response, "FE_documentation")
        elif output_format == "CSV":
            return formatted_response, save_as_csv(formatted_response, "FE_documentation")
        else:
            return formatted_response, save_as_txt(formatted_response, "FE_documentation")

    except Exception as e:
        error_msg = f"Error generating frontend validation: {str(e)}"
        temp_dir = tempfile.gettempdir()
        error_file = os.path.join(temp_dir, "error.txt")
        with open(error_file, 'w') as f:
            f.write(error_msg)
        return error_msg, error_file

if __name__ == "__main__":
    iface = create_interface()
    iface.launch()